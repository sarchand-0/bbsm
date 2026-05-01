from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1A, 0x2D, 0x40)
ORANGE  = RGBColor(0xE0, 0x78, 0x30)
GRAY    = RGBColor(0x55, 0x65, 0x7A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=ORANGE)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=NAVY)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=NAVY)
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=11, italic=italic, color=color or GRAY)
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f'"{text}"')
    set_font(run, size=11, italic=True, color=NAVY)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, color=GRAY)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11, color=GRAY)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 72)
    set_font(run, size=9, color=RGBColor(0xCC, 0xCC, 0xCC))
    return p

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
r = cover.add_run("BBSM ONLINE")
set_font(r, size=32, bold=True, color=ORANGE)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Executive Pitch & Live Demo Script")
set_font(r2, size=16, bold=False, color=NAVY)

aud = doc.add_paragraph()
aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
aud.paragraph_format.space_before = Pt(8)
r3 = aud.add_run("Prepared for the Board of Bhat-Bhateni Supermarket & Departmental Store")
set_font(r3, size=11, italic=True, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. THE HOOK
# ══════════════════════════════════════════════════════════════════════════════
heading1("THE HOOK — Open With This")
divider()

quote(
    "I want you to think about a customer. Her name is Sunita. She lives in Lazimpat. "
    "She has a Bhat-Bhateni Club Card — she's been a loyal customer for eleven years. "
    "Last Tuesday she needed cooking oil, rice, baby wipes, and shampoo. It was 6 PM. "
    "Traffic on Maharajgunj was locked. She opened her phone, went to Daraz, and bought "
    "all four items — things Bhat-Bhateni sells, things she trusts Bhat-Bhateni for — "
    "and gave that money to someone else. Not because she wanted to. "
    "But because Bhat-Bhateni wasn't there when she was ready to buy."
)

body("")
quote("She didn't abandon Bhat-Bhateni. Bhat-Bhateni wasn't available.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Pause. Let that land.")
set_font(r, size=10, italic=True, color=RGBColor(0x99, 0x99, 0x99))

# ══════════════════════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
heading1("THE PROBLEM — Make Them Feel It")
divider()

heading2("The problem is not competition. The problem is absence.")

body(
    "Right now, Bhat-Bhateni makes money in one way: a customer walks through a door. "
    "That is your entire revenue model. And it is brilliant — 27 stores, 22,000 employees, "
    "NPR 5.5 crore in daily sales. You built the most trusted retail brand in Nepal."
)
body("But here is what is happening silently:")

bullet("56% of Nepal is now online. 16.6 million internet users. 96% of them on mobile.  (DataReportal, 2025)")
bullet("Nepal's digital commerce is growing at 28.6% per year and will reach $3.94 billion by 2029.  (Statista)")
bullet("Daraz's grocery arm — dMart — is receiving 75,000 visits a day. Those are Bhat-Bhateni customers buying Bhat-Bhateni products from someone else.")
bullet("Every day that passes, a habit is forming. And habits are the hardest thing to break in retail.")

body("")
body("The window is still open. But it will not stay open.", color=ORANGE)

# ══════════════════════════════════════════════════════════════════════════════
# 3. THE OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════════
heading1("THE OPPORTUNITY — Numbers That Matter to a Board")
divider()

heading2("1.  New Revenue That Does Not Exist Today")
body(
    "People do not stop wanting things. They just stop going out to get them. "
    "Every person who cannot drive, will not brave Kathmandu traffic, or simply prefers their couch — "
    "that is a transaction Bhat-Bhateni is leaving behind. This is new revenue, "
    "not cannibalisation of your stores."
)
body(
    "Consider this: Nepali consumers already use online carts as shopping lists. "
    "Every time a new need arises, they open an app and add it. Whoever owns that app "
    "owns the purchase. Right now, that app is not Bhat-Bhateni's."
)

heading2("2.  Basket Size Goes Up Automatically")
body(
    "Walmart and Amazon consistently see 20–40% higher average order value online vs. in-store. "
    "Why? Because online shoppers browse without time pressure, receive recommendations, "
    "and add to cart the way people add to a wishlist. No rushing past an aisle "
    "because the children are pulling your arm."
)

heading2("3.  Geographic Reach Without a Single Brick")
body(
    "Bhat-Bhateni is in Pokhara, Chitwan, Dharan, Butwal. But the person 8 kilometres "
    "outside your Butwal branch — they cannot come to you. Online delivery brings them in. "
    "Every kilometre of delivery range is a kilometre of new market, at zero capital expenditure."
)

heading2("4.  The World Is Digitalising — and So Is Nepal's Competition")
body(
    "To maintain the monopoly Bhat-Bhateni has built over four decades, "
    "the brand must exist where the next generation of shoppers lives: on their phones. "
    "The cost of not being there is not zero. It is compounding daily."
)

heading2("5.  Data — The Asset No One Is Talking About")
body(
    "Every online transaction is structured data. What people search, what they add and "
    "remove from cart, what time they order, what neighbourhood they are in. "
    "This is the foundation of demand forecasting, inventory optimisation, "
    "and personalised promotions."
)
body(
    "As a data scientist who has built forecasting models for proof-of-concept projects, "
    "I can tell you — this data alone is worth as much as the revenue it generates. "
    "AI-powered inventory forecasting, automated replenishment triggers, "
    "personalised promotions by purchase history — these are not future concepts. "
    "They are implementable today with the right data pipeline."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. BEFORE vs AFTER TABLE
# ══════════════════════════════════════════════════════════════════════════════
heading1("BEFORE vs. AFTER")
divider()

table = doc.add_table(rows=8, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["", "Before", "After"]
rows_data = [
    ["Revenue hours",       "Store opening hours only",          "24 hours a day, 7 days a week"],
    ["Customers reached",   "People who can physically come",    "Anyone with a smartphone"],
    ["Purchase trigger",    "Customer remembers + travels",      "Customer sees → taps → orders"],
    ["Basket insight",      "Unknown",                           "Every item, time, location tracked"],
    ["Competitor exposure", "Limited — Bhat-Bhateni is everywhere", "Protected — habit formed online"],
    ["Promotions",          "Banners in-store",                  "Targeted push, by behaviour"],
    ["Geographic reach",    "Branch catchment area only",        "City-wide + surrounding districts"],
]

# Header row
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    shade_cell(cell, "1A2D40")
    p2 = cell.paragraphs[0]
    run = p2.add_run(h)
    set_font(run, size=10, bold=True, color=WHITE)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
for r_idx, row_data in enumerate(rows_data):
    row = table.rows[r_idx + 1]
    fill = "FFF8F0" if r_idx % 2 == 0 else "FFFFFF"
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        shade_cell(cell, fill)
        p2 = cell.paragraphs[0]
        run = p2.add_run(text)
        bold = c_idx == 0
        color = NAVY if c_idx == 0 else (ORANGE if c_idx == 2 else GRAY)
        set_font(run, size=10, bold=bold, color=color)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading1("FEATURES — What We Built")
divider()

heading2("Storefront — Customer Side")
bullet("8 departments: Groceries, Fresh Produce, Dairy, Beverages, Snacks, Household, Personal Care, Baby")
bullet("Smart search and filters — by price, availability, category, featured")
bullet("Product detail pages with live stock status")
bullet("Cart and checkout with saved address selection")
bullet("Discount codes and promotional pricing at checkout")
bullet("Order history and real-time delivery tracking with rider GPS")
bullet("Homepage promotions — Dashain Dhamaka, Weekly Deals, Department campaigns")
bullet("Find a Store — map with all branches, directions, nearest branch by location")

heading2("CRM — Staff and Management Side")
bullet("Live dashboard: today's orders, revenue, new customers, low-stock alerts")
bullet("Full order management: confirm → pack → ship → deliver in one click")
bullet("Product and inventory management: add, edit, archive, bulk CSV upload")
bullet("Customer profiles: order history, lifetime value, segmentation")
bullet("Discount codes and promotions: run a Dashain sale from your office in 30 seconds")
bullet("Sales reports: revenue by date range, top products, top customers, CSV download")

heading2("Delivery Operations")
bullet("Rider mobile app: assigned orders, GPS tracking, OTP delivery confirmation")
bullet("Admin live map: see all active riders, assign deliveries, monitor fleet")

# ══════════════════════════════════════════════════════════════════════════════
# 6. NOTE ON BUSINESS MODEL
# ══════════════════════════════════════════════════════════════════════════════
heading1("A Note on Your Business Model")
divider()

body("We know Bhat-Bhateni. We grew up going there.", italic=True)
body(
    "We know you have branches — not one store but a network of 27 across Nepal. "
    "The platform is built multi-location ready; branch-specific inventory management "
    "is on the roadmap and can be activated when you are ready to expand beyond Kathmandu."
)
body(
    "We know you have the Club Card. Points, tiers, member pricing — these are not small things. "
    "They are the reason Sunita has been coming for eleven years. "
    "The loyalty engine — points accumulation online, tier-based discounts, "
    "member-exclusive deals — is the next phase. What you are seeing today is the foundation. "
    "Solid, live, and ready to build on."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. DEMO SCRIPT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("LIVE DEMO SCRIPT")
divider()

p = doc.add_paragraph()
r = p.add_run("Total estimated time: 12–15 minutes  |  URLs: ")
set_font(r, size=10, italic=True, color=GRAY)
r2 = p.add_run("https://bbsm.vercel.app  (storefront)   https://bbsm.vercel.app/crm  (admin)")
set_font(r2, size=10, bold=True, color=ORANGE)
p.paragraph_format.space_after = Pt(12)

# Scene 1
heading2("Scene 1 — The Customer Experience  (2 min)")
quote("Let me be Sunita.")
numbered("Open https://bbsm.vercel.app — show the homepage: hero banner, departments, weekly deals, featured products.")
numbered("Click Fresh Produce — show the category page with filter sidebar.")
numbered('Search for "rice" — show search results, apply a price filter.')
numbered("Open a product detail page — show stock badge, images, price, Add to Cart.")
numbered("Add two more items from different departments.")
numbered("Click the cart icon — show the cart drawer with running total.")
numbered("Proceed to Checkout — select a saved address, enter a discount code, place the order.")
numbered("Show the Order Confirmation and live tracking page.")
quote("From browsing to confirmed order — under two minutes. No traffic. No parking. No queue.")

# Scene 2
heading2("Scene 2 — The Admin Sees It Instantly  (1.5 min)")
quote("Now let me be the manager at head office.")
numbered("Open https://bbsm.vercel.app/crm in a new tab.")
numbered("Login: admin@bbsm.np  /  admin123")
numbered("Show the Dashboard — the order we just placed is live. Revenue counter updated.")
numbered("Go to Orders → find the new order → click Advance Status → Confirmed.")
numbered("Go to Customers → open the customer profile → show lifetime value, order history.")
quote("Every order visible, every customer known, every rupee tracked — in real time.")

# Scene 3
heading2("Scene 3 — Run a Promotion in 30 Seconds  (1 min)")
numbered("Go to Promotions in the CRM — show the Dashain Dhamaka banner.")
numbered("Go to Discounts — show an active discount code.")
numbered("Go to Reports — show the revenue chart, top products, click Download CSV.")
quote("Your merchandising team runs a Dashain campaign. Your finance team downloads the revenue report. No developer. No IT ticket. Done.")

# Close
heading2("Closing Statement")
divider()
quote(
    "What you see today is not the finished product. It is proof that this is buildable, "
    "that it works, and that it works for how Bhat-Bhateni actually operates — "
    "multiple branches, loyal members, trusted products."
)
quote(
    "Nepal's digital commerce is growing at 28% a year. Daraz is already at your customer's door. "
    "The question is not whether Bhat-Bhateni should be online. "
    "The question is whether Bhat-Bhateni wants to own that space — or cede it."
)
quote(
    "You built Nepal's largest supermarket one store at a time. "
    "This is how you build the next one — without laying a single brick."
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\sarth\OneDrive\Desktop\BBSM_Executive_Pitch.docx"
doc.save(out)
print(f"Saved: {out}")